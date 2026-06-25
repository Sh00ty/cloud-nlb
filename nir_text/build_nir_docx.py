#!/usr/bin/env python3
"""Merge nir_text/*.md files into НИР_БАЛАНСЕР_ОТЧЕТ.docx.

Rules:
- Keep paragraphs 0..KEEP_LAST (title page + task assignment).
- Remove all old content beyond that boundary.
- Append markdown content with proper Word styles:
  * # / ## / ### / #### / ##### -> Heading 1..5
  * intro.md plain "ВВЕДЕНИЕ" line -> Heading 2
  * "- " / "* " lines -> list paragraph starting with em-dash
  * "1. " etc. numbered lists in references.md -> normal paragraphs (keep numbering inline)
  * **bold** spans -> bold runs
  * Image references ([Рисунок ...], "Снимок экрана ...") -> italic placeholder
- All new runs in Times New Roman, 14 pt.
- Lists use em-dash bullet "— ".
"""
from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

NIR_DIR = Path(__file__).resolve().parent
DOC_PATH = NIR_DIR / "НИР_БАЛАНСЕР_ОТЧЕТ.docx"
BACKUP_PATH = NIR_DIR / "НИР_БАЛАНСЕР_ОТЧЕТ.backup.docx"

# Files emitted before the TOC field. Per МИФИ / ГОСТ 7.32 the РЕФЕРАТ goes
# right after the task assignment and before СОДЕРЖАНИЕ; the TOC itself is then
# inserted programmatically, and ОПРЕДЕЛЕНИЯ + ВВЕДЕНИЕ follow.
FILES_BEFORE_TOC = [
    "abstract.md",
]

FILES_AFTER_TOC = [
    "definitions.md",
    "intro.md",
    "section_1_0_overview.md",
    "section_1_1_maglev.md",
    "section_1_2_unimog.md",
    "section_1_3_yandex.md",
    "section_1_4_requirements.md",
    "section_1_5_conclusion.md",
    "section_2_architecture.md",
    "section_3_dataplane.md",
    "section_4_agent.md",
    "section_5_cpl.md",
    "section_6_hc.md",
    "section_8_testing.md",
    "conclusion.md",
    "references.md",
]

# Last paragraph index of the "first pages" (inclusive) — task assignment ends at idx 97.
KEEP_LAST = 97

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)


def set_run_font(run):
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)
    run.font.size = FONT_SIZE


# Regex for **bold** spans.
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def clean_academic(text: str) -> str:
    """Strip markdown code spans (backticks) and replace arrows.

    Backticks are not allowed in academic text — keep the inner text without them.
    Arrows are replaced with em-dashes (also matches the work style).
    """
    text = text.replace("`", "")
    # Various arrow forms.
    text = text.replace("→", "—")
    text = text.replace("->", "—")
    return text


def add_styled_runs(paragraph, text: str, base_bold: bool = False, italic: bool = False) -> None:
    """Append runs to paragraph, parsing **bold** markup. Always TNR 14."""
    text = clean_academic(text)
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.bold = base_bold
            run.italic = italic
            set_run_font(run)
        run = paragraph.add_run(m.group(1))
        run.bold = True
        run.italic = italic
        set_run_font(run)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold
        run.italic = italic
        set_run_font(run)


def set_para_format(p, line_spacing: float = 1.5, first_line_indent_cm: float | None = 1.25,
                    space_before=0, space_after=0, justify: bool = True):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    if justify:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(doc, text: str, level: int):
    # All top-level chapter headings (numbered and non-numbered alike) share the
    # same Heading 1 style so Word's TOC field with `\o "1-3"` reliably picks
    # them all up. Non-numbered chapters (РЕФЕРАТ, СОДЕРЖАНИЕ, ОПРЕДЕЛЕНИЯ,
    # ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ) are detected by
    # the absence of a leading digit; they get centered alignment and zero
    # indent applied directly to the paragraph, while numbered ones stay
    # left-aligned with a first-line indent.
    centered = False
    if level == 1 and not re.match(r"^\d", text.strip()):
        centered = True
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    italic = level == 3 or level == 4
    add_styled_runs(p, text, base_bold=True, italic=italic)
    for run in p.runs:
        run.bold = True
        if italic:
            run.italic = True
    pf = p.paragraph_format
    if centered:
        pf.first_line_indent = Cm(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    # Explicit paragraph-level outlineLvl. Word's TOC field with \o picks up
    # paragraphs whose own pPr contains <w:outlineLvl>; relying on style
    # inheritance alone was unreliable in Word for Mac (numbered chapters
    # appeared but the unnumbered ones did not).
    spec = HEADING_STYLE_SPECS[level]
    pPr = p._element.get_or_add_pPr()
    for existing in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(existing)
    pPr.append(_new("w:outlineLvl", {"w:val": str(spec["outline"])}))
    # Page break before is normally configured on the style, but the centered
    # chapter headings need a forced break too (they used to live on Heading 5,
    # which had pageBreakBefore in its style).
    if centered:
        for existing in pPr.findall(qn("w:pageBreakBefore")):
            pPr.remove(existing)
        pPr.append(_new("w:pageBreakBefore", {"w:val": "1"}))
    return p


HEADING_STYLE_SPECS = {
    # level -> (size_pt, bold, italic, page_break_before, alignment, indent_twips, outline)
    1: dict(size=14, bold=True, italic=False, page_break=True,  align="left",   indent=709, outline=0),
    2: dict(size=14, bold=True, italic=False, page_break=False, align="left",   indent=709, outline=1),
    3: dict(size=14, bold=True, italic=True,  page_break=False, align="left",   indent=709, outline=2),
    4: dict(size=14, bold=False, italic=True, page_break=False, align="left",   indent=709, outline=3),
    # Heading 5 is no longer used for chapter headings, but kept for completeness.
    5: dict(size=14, bold=True, italic=False, page_break=True,  align="center", indent=0,   outline=0),
}


def _new(tag: str, attrs: dict | None = None):
    el = OxmlElement(tag)
    for k, v in (attrs or {}).items():
        el.set(qn(k), v)
    return el


def reset_heading_style(doc, level: int) -> None:
    spec = HEADING_STYLE_SPECS[level]
    style = doc.styles[f"Heading {level}"]
    el = style.element
    # Drop existing pPr/rPr if present.
    for tag in ("w:pPr", "w:rPr"):
        for child in list(el.findall(qn(tag))):
            el.remove(child)

    pPr = _new("w:pPr")
    pPr.append(_new("w:keepNext", {"w:val": "1"}))
    pPr.append(_new("w:keepLines", {"w:val": "1"}))
    if spec["page_break"]:
        pPr.append(_new("w:pageBreakBefore", {"w:val": "1"}))
    pPr.append(_new("w:spacing", {"w:before": "240", "w:after": "120", "w:lineRule": "auto"}))
    pPr.append(_new("w:ind", {"w:firstLine": str(spec["indent"])}))
    pPr.append(_new("w:jc", {"w:val": spec["align"]}))
    pPr.append(_new("w:outlineLvl", {"w:val": str(spec["outline"])}))

    rPr = _new("w:rPr")
    rFonts = _new("w:rFonts", {
        "w:ascii": FONT_NAME, "w:hAnsi": FONT_NAME,
        "w:cs": FONT_NAME, "w:eastAsia": FONT_NAME,
    })
    rPr.append(rFonts)
    if spec["bold"]:
        rPr.append(_new("w:b", {"w:val": "1"}))
        rPr.append(_new("w:bCs", {"w:val": "1"}))
    else:
        rPr.append(_new("w:b", {"w:val": "0"}))
        rPr.append(_new("w:bCs", {"w:val": "0"}))
    if spec["italic"]:
        rPr.append(_new("w:i", {"w:val": "1"}))
        rPr.append(_new("w:iCs", {"w:val": "1"}))
    else:
        rPr.append(_new("w:i", {"w:val": "0"}))
        rPr.append(_new("w:iCs", {"w:val": "0"}))
    sz = str(spec["size"] * 2)  # half-points
    rPr.append(_new("w:sz", {"w:val": sz}))
    rPr.append(_new("w:szCs", {"w:val": sz}))
    rPr.append(_new("w:color", {"w:val": "000000"}))

    # Maintain canonical order: name/basedOn/next come first; pPr after them.
    name = el.find(qn("w:name"))
    insert_after = name
    for tag in ("w:basedOn", "w:next", "w:link", "w:uiPriority", "w:qFormat"):
        node = el.find(qn(tag))
        if node is not None:
            insert_after = node
    insert_after.addnext(pPr)
    pPr.addnext(rPr)


def fix_heading_styles(doc) -> None:
    for level in HEADING_STYLE_SPECS:
        try:
            reset_heading_style(doc, level)
        except KeyError:
            pass


def add_body_para(doc, text: str, italic: bool = False, indent: bool = True,
                   center: bool = False):
    p = doc.add_paragraph()
    p.style = doc.styles["normal"]
    add_styled_runs(p, text, italic=italic)
    set_para_format(p, first_line_indent_cm=1.25 if indent else 0, justify=not center)
    if center:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


IMAGE_SEARCH_DIRS = ("", "nlb_nir")


def resolve_image_path(name: str) -> Path | None:
    """Resolve an image filename against NIR_DIR and NIR_DIR/nlb_nir/.

    Returns the first existing path or ``None`` if nothing matches. Bare
    filenames and relative paths are both accepted.
    """
    candidate = Path(name)
    if candidate.is_absolute() and candidate.exists():
        return candidate
    for sub in IMAGE_SEARCH_DIRS:
        p = NIR_DIR / sub / name
        if p.exists():
            return p
    return None


def _figure_caption_para(doc, caption: str, number: int):
    cap = doc.add_paragraph()
    cap.style = doc.styles["normal"]
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    cap.paragraph_format.line_spacing = 1.5
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(12)
    add_styled_runs(cap, f"Рисунок {number} — {caption}")


def add_figure_with_image(doc, image_path: Path, caption: str, number: int,
                           width_cm: float = 15.0) -> None:
    """Embed an image, centered, with caption «Рисунок N — Caption» beneath."""
    p = doc.add_paragraph()
    p.style = doc.styles["normal"]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    set_run_font(run)
    run.add_picture(str(image_path), width=Cm(width_cm))
    _figure_caption_para(doc, caption, number)


def add_figure_placeholder(doc, caption: str, number: int) -> None:
    """Placeholder paragraph + centered caption «Рисунок N — Название»."""
    placeholder = doc.add_paragraph()
    placeholder.style = doc.styles["normal"]
    placeholder.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder.paragraph_format.first_line_indent = Cm(0)
    placeholder.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    placeholder.paragraph_format.line_spacing = 1.5
    placeholder.paragraph_format.space_before = Pt(6)
    placeholder.paragraph_format.space_after = Pt(0)
    run = placeholder.add_run("[место для рисунка]")
    run.italic = True
    set_run_font(run)

    _figure_caption_para(doc, caption, number)


def add_list_item(doc, text: str):
    """List item rendered as paragraph starting with em-dash."""
    p = doc.add_paragraph()
    p.style = doc.styles["normal"]
    add_styled_runs(p, "— " + text)
    set_para_format(p, first_line_indent_cm=1.25)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.style = doc.styles["normal"]
    run = p.add_run()
    set_run_font(run)
    run.add_break(WD_BREAK.PAGE)
    return p


HEADING_RE = re.compile(r"^(#{1,5})\s+(.*)$")
LIST_RE = re.compile(r"^[-*]\s+(.*)$")
IMG_BRACKET_RE = re.compile(r"^\[Рисунок[^]]*\]\s*$")
IMG_BRACKET_PATH_RE = re.compile(r"^\[Рисунок[^]]*\]\(([^)]+)\)\s*$")
IMG_MD_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
SCREENSHOT_RE = re.compile(r"^(Снимок экрана.+?\.png)\s*—\s*(.*)$")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|\s*$")


def _split_table_row(row: str) -> list[str]:
    inner = row.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


def _set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def add_md_table(doc, header: list[str], rows: list[list[str]]) -> None:
    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.autofit = True
    # Header row.
    for j, txt in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        add_styled_runs(p, txt, base_bold=True)
        for run in p.runs:
            run.bold = True
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        _set_cell_borders(cell)
    # Body rows.
    for i, row in enumerate(rows, start=1):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            value = row[j] if j < len(row) else ""
            add_styled_runs(p, value)
            pf = p.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.15
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            _set_cell_borders(cell)


FIG_COUNTER = {"n": 0}


def _add_borderless_cell_with_text(cell, text: str, bold_lead: bool = False) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    p.style = cell.part.document.styles["normal"] if False else p.style
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    add_styled_runs(p, text, base_bold=bold_lead)
    # Set borders to "nil" for invisible table.
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tc_pr.append(tcBorders)


DEF_LINE_RE = re.compile(r"^\*\*([^*]+)\*\*\s*—\s*(.+)$")


def render_definitions_table(doc, lines: list[str]) -> None:
    """Render a `**Term** — Definition` list as a borderless two-column table."""
    pairs: list[tuple[str, str]] = []
    for line in lines:
        s = line.strip()
        m = DEF_LINE_RE.match(s)
        if m:
            pairs.append((m.group(1).strip(), m.group(2).strip()))
    if not pairs:
        return
    table = doc.add_table(rows=len(pairs), cols=2)
    table.autofit = False
    # 30% / 70% split of page content (page width ~16cm after margins).
    total_dxa = 9000
    left_dxa, right_dxa = 2700, 6300
    for ri, (term, defn) in enumerate(pairs):
        row = table.rows[ri]
        for j, w in enumerate((left_dxa, right_dxa)):
            row.cells[j].width = Cm(w / 567.0)  # 567 DXA = 1 cm
        _add_borderless_cell_with_text(row.cells[0], term, bold_lead=False)
        _add_borderless_cell_with_text(row.cells[1], "— " + defn, bold_lead=False)


def process_md_file(doc, path: Path) -> None:
    raw = path.read_text(encoding="utf-8")
    lines = raw.split("\n")
    # Trim trailing empty lines
    while lines and lines[-1].strip() == "":
        lines.pop()

    name = path.name

    # Special handling: definitions.md → render the term/definition list as a borderless table.
    if name == "definitions.md":
        # Emit heading + intro paragraph, then table from the **term** — definition lines.
        heading_done = False
        intro_done = False
        def_lines: list[str] = []
        for line in lines:
            s = line.strip()
            if not s:
                continue
            if not heading_done and s.startswith("##"):
                m = HEADING_RE.match(s)
                if m:
                    add_heading(doc, m.group(2).strip(), max(len(m.group(1)) - 1, 1))
                    heading_done = True
                    continue
            if DEF_LINE_RE.match(s):
                def_lines.append(s)
            elif not intro_done:
                add_body_para(doc, s)
                intro_done = True
        render_definitions_table(doc, def_lines)
        return

    # Special: intro.md begins with a bare "ВВЕДЕНИЕ" heading.
    if name == "intro.md":
        idx = 0
        while idx < len(lines) and lines[idx].strip() == "":
            idx += 1
        if idx < len(lines) and lines[idx].strip().upper() == "ВВЕДЕНИЕ":
            add_heading(doc, "ВВЕДЕНИЕ", 1)
            lines = lines[idx + 1:]

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        # Heading?
        m = HEADING_RE.match(stripped)
        if m:
            hashes, text = m.group(1), m.group(2).strip()
            md_level = len(hashes)
            # md '##' -> Word Heading 1 (top-level chapters with page break);
            # md '###' -> Heading 2; '####' -> Heading 3; etc.
            word_level = max(md_level - 1, 1)
            add_heading(doc, text, word_level)
            i += 1
            continue

        # Markdown table?
        if TABLE_ROW_RE.match(stripped):
            # Look ahead for separator row.
            if i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1].strip()):
                header_cells = _split_table_row(stripped)
                j = i + 2
                body_rows: list[list[str]] = []
                while j < len(lines) and TABLE_ROW_RE.match(lines[j].strip()):
                    body_rows.append(_split_table_row(lines[j].strip()))
                    j += 1
                add_md_table(doc, header_cells, body_rows)
                i = j
                continue

        # List item: collect contiguous lines.
        m = LIST_RE.match(stripped)
        if m:
            add_list_item(doc, m.group(1).strip())
            i += 1
            continue

        # Markdown image syntax «![Подпись](path)» — embed the image and add
        # an auto-numbered caption underneath.
        mi = IMG_MD_RE.match(stripped)
        if mi:
            FIG_COUNTER["n"] += 1
            caption = mi.group(1).strip()
            img_path = resolve_image_path(mi.group(2).strip())
            if img_path is not None:
                add_figure_with_image(doc, img_path, caption, FIG_COUNTER["n"])
            else:
                print(f"    ! image not found: {mi.group(2)} (рисунок {FIG_COUNTER['n']})",
                      file=sys.stderr)
                add_figure_placeholder(doc, caption, FIG_COUNTER["n"])
            i += 1
            continue

        # Bracketed placeholder with explicit path: «[Рисунок N — caption](path)»
        bp = IMG_BRACKET_PATH_RE.match(stripped)
        if bp:
            FIG_COUNTER["n"] += 1
            inner = stripped[1:stripped.index("]")].strip()
            cap = re.sub(r"^Рисунок\s*\S+\s*—\s*", "", inner)
            img_path = resolve_image_path(bp.group(1).strip())
            if img_path is not None:
                add_figure_with_image(doc, img_path, cap, FIG_COUNTER["n"])
            else:
                print(f"    ! image not found: {bp.group(1)} (рисунок {FIG_COUNTER['n']})",
                      file=sys.stderr)
                add_figure_placeholder(doc, cap, FIG_COUNTER["n"])
            i += 1
            continue

        # Bracketed placeholder without a path (kept for backward compatibility).
        if IMG_BRACKET_RE.match(stripped):
            FIG_COUNTER["n"] += 1
            inner = stripped[1:-1].strip()
            # Drop leading "Рисунок N — " if present in source.
            cap = re.sub(r"^Рисунок\s*\S+\s*—\s*", "", inner)
            add_figure_placeholder(doc, cap, FIG_COUNTER["n"])
            i += 1
            continue

        # Screenshot-style reference: «Снимок экрана … .png — caption». If the
        # PNG exists under nir_text/ or nir_text/nlb_nir/, embed it; otherwise
        # fall back to a placeholder.
        sm = SCREENSHOT_RE.match(stripped)
        if sm:
            FIG_COUNTER["n"] += 1
            filename = sm.group(1).strip()
            caption = sm.group(2).strip()
            img_path = resolve_image_path(filename)
            if img_path is not None:
                add_figure_with_image(doc, img_path, caption, FIG_COUNTER["n"])
            else:
                print(f"    ! screenshot not found: {filename} (рисунок {FIG_COUNTER['n']})",
                      file=sys.stderr)
                add_figure_placeholder(doc, caption, FIG_COUNTER["n"])
            i += 1
            continue

        # Otherwise plain paragraph (treat each non-empty line as its own paragraph).
        add_body_para(doc, stripped)
        i += 1


def add_toc(doc) -> None:
    """Insert a Word-native TOC field.

    ``w:dirty="true"`` on the field start marks the TOC as needing a refresh,
    so Word repopulates entries automatically on open (without the user
    pressing F9). All chapter headings — numbered (1, 2, …) and unnumbered
    (РЕФЕРАТ, ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, …) — use Heading 1 with explicit
    paragraph-level ``<w:outlineLvl w:val="0"/>``, so ``\\o "1-3"`` is enough
    to capture all of them at TOC level 1; subsection headings (Heading 2)
    land at TOC level 2, and ``\\u`` instructs Word to honour paragraph-level
    outline overrides even when they differ from the style.
    """
    add_heading(doc, "СОДЕРЖАНИЕ", 1)

    holder = doc.add_paragraph()
    holder.style = doc.styles["normal"]
    holder.paragraph_format.first_line_indent = Cm(0)
    run = holder.add_run()
    set_run_font(run)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    fldChar1.set(qn("w:dirty"), "true")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Содержание сформируется при открытии в Word."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(placeholder)
    run._r.append(fldChar4)


def enable_update_fields(doc) -> None:
    """Set ``<w:updateFields w:val="true"/>`` in settings.xml.

    Without this flag Word warns the user before refreshing field contents on
    open. With it set, the dirty TOC is regenerated automatically.
    """
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.insert(0, upd)


def set_page_setup(doc) -> None:
    """Enforce ГОСТ/МФТИ page geometry on every section.

    Left margin 30 mm, the other three 20 mm — per the МФТИ requirement, which
    is stricter than ГОСТ 7.32's 15 mm right margin. ``different_first_page``
    keeps the page number off the first sheet (ГОСТ: the title sheet is counted
    in the numbering but its number is not printed). Note: this suppresses the
    number only on the very first page; if the front matter spans several pages,
    revisit once the final title layout is fixed.
    """
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.different_first_page_header_footer = True


def add_page_numbers(doc) -> None:
    """Center page number in the footer of the default section."""
    section = doc.sections[0]
    footer = section.footer
    # Reuse the first existing paragraph, or create one.
    if footer.paragraphs:
        p = footer.paragraphs[0]
        # Clear existing runs.
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
    else:
        p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)

    run = p.add_run()
    set_run_font(run)
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(placeholder)
    run._r.append(fldEnd)


def strip_old_content(doc, keep_last_idx: int) -> None:
    """Remove everything after the task assignment.

    The backup file ends the task assignment with the signature table containing
    «Руководитель ВКР». All body siblings after that table (old bachelor's TOC,
    definitions table, chapter content) must be removed before appending the new
    content. ``doc.paragraphs`` only enumerates top-level paragraphs, so an older
    paragraph-index-based strip leaves the bachelor's tables behind — hence the
    explicit walk over ``<w:body>`` children below. The final ``<w:sectPr>`` is
    preserved so the document remains valid.
    """
    body = doc.element.body
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w_t = f"{{{w_ns}}}t"
    children = list(body)

    cutoff = None
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1]
        if tag != "tbl":
            continue
        texts = [t.text or "" for t in child.iter(w_t)]
        if any("Руководитель ВКР" in t for t in texts):
            cutoff = i
            break

    if cutoff is None:
        # Fall back to the previous (paragraph-based) behaviour if the signature
        # table cannot be located — keeps the script working on hand-edited files.
        paragraphs = doc.paragraphs
        to_remove = [p._element for p in paragraphs[keep_last_idx + 1:]]
        for el in to_remove:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        return

    for child in children[cutoff + 1:]:
        tag = child.tag.split("}")[-1]
        if tag == "sectPr":
            continue
        body.remove(child)


def main():
    # Start from the pristine backup every time so the script is idempotent
    # (otherwise stripping "old content" would eat into the previously generated body).
    if BACKUP_PATH.exists():
        import shutil
        shutil.copyfile(BACKUP_PATH, DOC_PATH)
        print(f"Restored {DOC_PATH.name} from backup.")
    else:
        print(f"WARNING: backup not found at {BACKUP_PATH}; using {DOC_PATH} as-is.",
              file=sys.stderr)
    doc = Document(str(DOC_PATH))
    print(f"Loaded doc: {len(doc.paragraphs)} paragraphs total")
    strip_old_content(doc, KEEP_LAST)
    print(f"After strip: {len(doc.paragraphs)} paragraphs kept")
    fix_heading_styles(doc)
    print("Heading styles normalized.")
    set_page_setup(doc)
    print("Page geometry set (left 30 mm, others 20 mm; first page unnumbered).")
    add_page_numbers(doc)
    print("Page numbers added to footer.")

    # МИФИ/ГОСТ order: РЕФЕРАТ before СОДЕРЖАНИЕ, then ОПРЕДЕЛЕНИЯ, then ВВЕДЕНИЕ.
    for fname in FILES_BEFORE_TOC:
        path = NIR_DIR / fname
        if not path.exists():
            print(f"  ! missing: {fname}", file=sys.stderr)
            continue
        print(f"  + {fname}")
        process_md_file(doc, path)

    add_toc(doc)
    print("TOC field inserted.")
    enable_update_fields(doc)
    print("updateFields enabled in settings.")

    for fname in FILES_AFTER_TOC:
        path = NIR_DIR / fname
        if not path.exists():
            print(f"  ! missing: {fname}", file=sys.stderr)
            continue
        print(f"  + {fname}")
        process_md_file(doc, path)

    out = DOC_PATH
    doc.save(str(out))
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
